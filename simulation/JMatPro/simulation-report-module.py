"""
材料仿真报告生成模块
用于生成JMatPro等材料仿真软件的专业报告
支持Word和Excel格式输出

依赖安装：
pip install python-docx openpyxl matplotlib pandas Pillow numpy

使用示例：
from simulation_report import SimulationReportGenerator

# 创建报告生成器
report = SimulationReportGenerator()

# 设置基本信息
report.set_basic_info(
    title="45号钢材料性能仿真分析报告",
    author="材料研究部",
    company="XX科技有限公司"
)

# 添加材料成分
report.add_composition({'C': 0.45, 'Mn': 1.2, 'Si': 0.3})

# 生成报告
report.generate_word_report("output/report.docx")
report.generate_excel_report("output/data.xlsx")
"""

import os
import datetime
from typing import Dict, List, Tuple, Optional, Any
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.chart import LineChart, Reference
from openpyxl.drawing.image import Image as XLImage
from PIL import Image
import io
import warnings

# 配置matplotlib中文显示
plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
warnings.filterwarnings('ignore')


class SimulationReportGenerator:
    """材料仿真报告生成器"""
    
    def __init__(self):
        """初始化报告生成器"""
        self.basic_info = {
            'title': '材料仿真分析报告',
            'subtitle': '',
            'author': '未指定',
            'company': '未指定',
            'date': datetime.datetime.now().strftime('%Y年%m月%d日'),
            'report_number': f"REP-{datetime.datetime.now().strftime('%Y%m%d-%H%M%S')}"
        }
        
        self.material_info = {
            'name': '',
            'type': '',
            'standard': '',
            'composition': {},
            'description': ''
        }
        
        self.simulation_params = {
            'software': 'JMatPro',
            'version': '',
            'modules': [],
            'temperature_range': [],
            'cooling_rates': [],
            'grain_size': 50,
            'other_params': {}
        }
        
        self.results = {
            'phase_diagram': None,
            'ttt_diagram': None,
            'cct_diagram': None,
            'solidification': None,
            'mechanical_properties': None,
            'thermal_properties': None,
            'custom_results': {}
        }
        
        self.figures = []
        self.tables = []
        self.conclusions = []
        
        # 创建临时图片目录
        self.temp_dir = 'temp_figures'
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
    
    def set_basic_info(self, **kwargs):
        """设置报告基本信息"""
        self.basic_info.update(kwargs)
    
    def set_material_info(self, **kwargs):
        """设置材料信息"""
        self.material_info.update(kwargs)
    
    def set_simulation_params(self, **kwargs):
        """设置仿真参数"""
        self.simulation_params.update(kwargs)
    
    def add_composition(self, composition: Dict[str, float]):
        """添加材料成分"""
        self.material_info['composition'] = composition
    
    def add_phase_diagram_results(self, temperatures: List[float], phases: Dict[str, List[float]]):
        """添加相图结果"""
        self.results['phase_diagram'] = {
            'temperatures': temperatures,
            'phases': phases
        }
        
        # 自动生成相图
        fig = self._create_phase_diagram(temperatures, phases)
        self.figures.append({
            'name': 'phase_diagram',
            'figure': fig,
            'caption': '平衡相图',
            'description': '不同温度下各相的体积分数变化'
        })
    
    def add_ttt_results(self, ttt_data: Dict[str, Any]):
        """添加TTT曲线结果"""
        self.results['ttt_diagram'] = ttt_data
        
        # 自动生成TTT图
        fig = self._create_ttt_diagram(ttt_data)
        self.figures.append({
            'name': 'ttt_diagram',
            'figure': fig,
            'caption': 'TTT转变曲线',
            'description': '时间-温度-转变图，显示不同相变的开始和结束时间'
        })
    
    def add_cct_results(self, cct_data: Dict[float, Any]):
        """添加CCT曲线结果"""
        self.results['cct_diagram'] = cct_data
        
        # 自动生成CCT图
        fig = self._create_cct_diagram(cct_data)
        self.figures.append({
            'name': 'cct_diagram',
            'figure': fig,
            'caption': 'CCT转变曲线',
            'description': '连续冷却转变图，不同冷却速率下的相变行为'
        })
    
    def add_mechanical_properties(self, properties: Dict[str, Any]):
        """添加机械性能数据"""
        self.results['mechanical_properties'] = properties
        
        # 生成机械性能图表
        if 'temperature' in properties:
            fig = self._create_mechanical_properties_chart(properties)
            self.figures.append({
                'name': 'mechanical_properties',
                'figure': fig,
                'caption': '机械性能随温度变化',
                'description': '屈服强度、抗拉强度、硬度等随温度的变化规律'
            })
    
    def add_solidification_results(self, solidification_data: Dict[str, Any]):
        """添加凝固模拟结果"""
        self.results['solidification'] = solidification_data
        
        # 生成凝固曲线图
        if 'temperature' in solidification_data and 'solid_fraction' in solidification_data:
            fig = self._create_solidification_curve(solidification_data)
            self.figures.append({
                'name': 'solidification_curve',
                'figure': fig,
                'caption': '凝固曲线',
                'description': '固相分数随温度的变化关系'
            })
    
    def add_custom_figure(self, figure: plt.Figure, name: str, caption: str, description: str = ''):
        """添加自定义图表"""
        self.figures.append({
            'name': name,
            'figure': figure,
            'caption': caption,
            'description': description
        })
    
    def add_custom_table(self, data: pd.DataFrame, name: str, caption: str):
        """添加自定义表格"""
        self.tables.append({
            'name': name,
            'data': data,
            'caption': caption
        })
    
    def add_conclusion(self, conclusion: str):
        """添加结论"""
        self.conclusions.append(conclusion)
    
    def _create_phase_diagram(self, temperatures: List[float], phases: Dict[str, List[float]]) -> plt.Figure:
        """创建相图"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FECA57', '#48C9B0']
        
        for i, (phase_name, fractions) in enumerate(phases.items()):
            color = colors[i % len(colors)]
            ax.plot(temperatures, fractions, label=phase_name, linewidth=2.5, color=color)
        
        ax.set_xlabel('温度 (°C)', fontsize=12)
        ax.set_ylabel('相分数', fontsize=12)
        ax.set_title('平衡相图', fontsize=14, fontweight='bold')
        ax.legend(loc='best', frameon=True, fancybox=True, shadow=True)
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.set_xlim(min(temperatures), max(temperatures))
        ax.set_ylim(0, 1.05)
        
        plt.tight_layout()
        return fig
    
    def _create_ttt_diagram(self, ttt_data: Dict[str, Any]) -> plt.Figure:
        """创建TTT图"""
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # 示例TTT曲线绘制
        phases = ['Pearlite', 'Bainite', 'Martensite']
        colors = ['red', 'blue', 'green']
        
        for phase, color in zip(phases, colors):
            if phase in ttt_data:
                start_times = ttt_data[phase].get('start_times', [])
                start_temps = ttt_data[phase].get('start_temperatures', [])
                finish_times = ttt_data[phase].get('finish_times', [])
                finish_temps = ttt_data[phase].get('finish_temperatures', [])
                
                if start_times and start_temps:
                    ax.semilogx(start_times, start_temps, color=color, linewidth=2, 
                               label=f'{phase} 开始', linestyle='--')
                if finish_times and finish_temps:
                    ax.semilogx(finish_times, finish_temps, color=color, linewidth=2,
                               label=f'{phase} 结束')
        
        ax.set_xlabel('时间 (秒)', fontsize=12)
        ax.set_ylabel('温度 (°C)', fontsize=12)
        ax.set_title('TTT转变曲线', fontsize=14, fontweight='bold')
        ax.legend(loc='best', frameon=True, fancybox=True, shadow=True)
        ax.grid(True, alpha=0.3, which='both', linestyle='--')
        ax.set_xlim(0.1, 10000)
        ax.set_ylim(200, 800)
        
        plt.tight_layout()
        return fig
    
    def _create_cct_diagram(self, cct_data: Dict[float, Any]) -> plt.Figure:
        """创建CCT图"""
        fig, ax = plt.subplots(figsize=(10, 8))
        
        colors = plt.cm.coolwarm(np.linspace(0, 1, len(cct_data)))
        
        for i, (cooling_rate, data) in enumerate(cct_data.items()):
            if 'temperature' in data and 'time' in data:
                ax.semilogx(data['time'], data['temperature'], 
                           color=colors[i], linewidth=2,
                           label=f'{cooling_rate} °C/s')
        
        ax.set_xlabel('时间 (秒)', fontsize=12)
        ax.set_ylabel('温度 (°C)', fontsize=12)
        ax.set_title('CCT连续冷却转变图', fontsize=14, fontweight='bold')
        ax.legend(loc='best', title='冷却速率', frameon=True, fancybox=True, shadow=True)
        ax.grid(True, alpha=0.3, which='both', linestyle='--')
        
        plt.tight_layout()
        return fig
    
    def _create_mechanical_properties_chart(self, properties: Dict[str, Any]) -> plt.Figure:
        """创建机械性能图表"""
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 10))
        
        temps = properties.get('temperature', [])
        
        # 强度图
        if 'yield_strength' in properties:
            ax1.plot(temps, properties['yield_strength'], 'o-', color='#FF6B6B', 
                    linewidth=2.5, markersize=8, label='屈服强度')
        if 'tensile_strength' in properties:
            ax1.plot(temps, properties['tensile_strength'], 's-', color='#4ECDC4',
                    linewidth=2.5, markersize=8, label='抗拉强度')
        
        ax1.set_xlabel('温度 (°C)', fontsize=12)
        ax1.set_ylabel('强度 (MPa)', fontsize=12)
        ax1.set_title('强度随温度变化', fontsize=13, fontweight='bold')
        ax1.legend(loc='best', frameon=True, fancybox=True, shadow=True)
        ax1.grid(True, alpha=0.3, linestyle='--')
        
        # 硬度图
        if 'hardness' in properties:
            ax2.plot(temps, properties['hardness'], '^-', color='#45B7D1',
                    linewidth=2.5, markersize=8, label='硬度')
            ax2.set_xlabel('温度 (°C)', fontsize=12)
            ax2.set_ylabel('硬度 (HV)', fontsize=12)
            ax2.set_title('硬度随温度变化', fontsize=13, fontweight='bold')
            ax2.legend(loc='best', frameon=True, fancybox=True, shadow=True)
            ax2.grid(True, alpha=0.3, linestyle='--')
        
        plt.tight_layout()
        return fig
    
    def _create_solidification_curve(self, solidification_data: Dict[str, Any]) -> plt.Figure:
        """创建凝固曲线"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        temps = solidification_data['temperature']
        solid_fraction = solidification_data['solid_fraction']
        
        ax.plot(temps, solid_fraction, 'b-', linewidth=3)
        ax.fill_between(temps, 0, solid_fraction, alpha=0.3, color='skyblue')
        
        ax.set_xlabel('温度 (°C)', fontsize=12)
        ax.set_ylabel('固相分数', fontsize=12)
        ax.set_title('凝固曲线', fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.set_xlim(max(temps), min(temps))
        ax.set_ylim(0, 1.05)
        
        # 添加特征点标注
        if 'liquidus_temp' in solidification_data:
            ax.axvline(solidification_data['liquidus_temp'], color='red', 
                      linestyle='--', alpha=0.7, label='液相线')
        if 'solidus_temp' in solidification_data:
            ax.axvline(solidification_data['solidus_temp'], color='green',
                      linestyle='--', alpha=0.7, label='固相线')
        
        if ax.get_legend_handles_labels()[0]:
            ax.legend(loc='best', frameon=True, fancybox=True, shadow=True)
        
        plt.tight_layout()
        return fig
    
    def _save_figure(self, fig: plt.Figure, name: str) -> str:
        """保存图表并返回路径"""
        filepath = os.path.join(self.temp_dir, f'{name}.png')
        fig.savefig(filepath, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return filepath
    
    def _add_header_footer(self, doc: Document):
        """添加页眉页脚"""
        section = doc.sections[0]
        
        # 页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"{self.basic_info['company']} - {self.basic_info['title']}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"报告编号: {self.basic_info['report_number']}    第 "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _create_cover_page(self, doc: Document):
        """创建封面页"""
        # 公司名称
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(48)
        run = p.add_run(self.basic_info['company'])
        run.font.size = Pt(24)
        run.font.bold = True
        
        # 标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(120)
        p.space_after = Pt(24)
        run = p.add_run(self.basic_info['title'])
        run.font.size = Pt(36)
        run.font.bold = True
        
        # 副标题
        if self.basic_info.get('subtitle'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_after = Pt(120)
            run = p.add_run(self.basic_info['subtitle'])
            run.font.size = Pt(18)
        
        # 材料信息
        if self.material_info.get('name'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(48)
            run = p.add_run(f"材料: {self.material_info['name']}")
            run.font.size = Pt(16)
        
        # 报告信息
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(180)
        p.add_run(f"作者: {self.basic_info['author']}\n").font.size = Pt(14)
        p.add_run(f"日期: {self.basic_info['date']}\n").font.size = Pt(14)
        p.add_run(f"报告编号: {self.basic_info['report_number']}").font.size = Pt(14)
        
        doc.add_page_break()
    
    def _create_abstract(self, doc: Document):
        """创建摘要页"""
        doc.add_heading('摘要', level=1)
        
        abstract_text = f"""
本报告采用{self.simulation_params['software']}软件对{self.material_info.get('name', '材料')}进行了全面的仿真分析。
主要计算模块包括: {', '.join(self.simulation_params.get('modules', ['相平衡', 'TTT/CCT', '凝固', '机械性能']))}。

材料成分:
"""
        doc.add_paragraph(abstract_text)
        
        # 成分表格
        if self.material_info['composition']:
            table = doc.add_table(rows=2, cols=len(self.material_info['composition']) + 1)
            table.style = 'Light Shading Accent 1'
            
            # 表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '元素'
            for i, element in enumerate(self.material_info['composition'].keys()):
                hdr_cells[i + 1].text = element
            
            # 数值
            row_cells = table.rows[1].cells
            row_cells[0].text = '含量 (%)'
            for i, value in enumerate(self.material_info['composition'].values()):
                row_cells[i + 1].text = f'{value:.2f}'
        
        # 关键发现
        doc.add_paragraph('\n关键发现:')
        if self.conclusions:
            for i, conclusion in enumerate(self.conclusions[:3]):  # 只显示前3条
                doc.add_paragraph(f'{i+1}. {conclusion}', style='List Bullet')
        
        doc.add_page_break()
    
    def _add_composition_section(self, doc: Document):
        """添加成分章节"""
        doc.add_heading('1. 材料成分', level=1)
        
        # 材料基本信息
        if self.material_info.get('type'):
            doc.add_paragraph(f"材料类型: {self.material_info['type']}")
        if self.material_info.get('standard'):
            doc.add_paragraph(f"执行标准: {self.material_info['standard']}")
        
        # 成分详细表格
        if self.material_info['composition']:
            doc.add_paragraph('\n化学成分分析:')
            
            # 创建更详细的成分表
            elements = list(self.material_info['composition'].keys())
            values = list(self.material_info['composition'].values())
            
            # 分两行显示，更美观
            half = (len(elements) + 1) // 2
            table = doc.add_table(rows=3, cols=half * 2)
            table.style = 'Medium Grid 3 Accent 1'
            
            # 填充表格
            for i in range(half):
                row_idx = i * 2
                if i < len(elements):
                    table.cell(0, row_idx).text = elements[i]
                    table.cell(1, row_idx).text = f'{values[i]:.3f}'
                    table.cell(2, row_idx).text = 'wt%'
                
                if i + half < len(elements):
                    table.cell(0, row_idx + 1).text = elements[i + half]
                    table.cell(1, row_idx + 1).text = f'{values[i + half]:.3f}'
                    table.cell(2, row_idx + 1).text = 'wt%'
        
        if self.material_info.get('description'):
            doc.add_paragraph(f"\n材料说明: {self.material_info['description']}")
    
    def _add_parameters_section(self, doc: Document):
        """添加仿真参数章节"""
        doc.add_heading('2. 仿真参数', level=1)
        
        doc.add_paragraph(f"仿真软件: {self.simulation_params['software']}")
        if self.simulation_params.get('version'):
            doc.add_paragraph(f"软件版本: {self.simulation_params['version']}")
        
        doc.add_paragraph('\n计算参数:')
        
        # 参数表格
        params = []
        if self.simulation_params.get('temperature_range'):
            params.append(('温度范围', f"{self.simulation_params['temperature_range'][0]} - {self.simulation_params['temperature_range'][1]} °C"))
        if self.simulation_params.get('cooling_rates'):
            params.append(('冷却速率', ', '.join([f"{rate} °C/s" for rate in self.simulation_params['cooling_rates']])))
        if self.simulation_params.get('grain_size'):
            params.append(('晶粒尺寸', f"{self.simulation_params['grain_size']} μm"))
        
        # 添加其他参数
        for key, value in self.simulation_params.get('other_params', {}).items():
            params.append((key, str(value)))
        
        if params:
            table = doc.add_table(rows=len(params) + 1, cols=2)
            table.style = 'Light List Accent 1'
            
            # 表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '参数名称'
            hdr_cells[1].text = '参数值'
            
            # 参数值
            for i, (param_name, param_value) in enumerate(params):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = param_name
                row_cells[1].text = param_value
    
    def _add_results_section(self, doc: Document):
        """添加结果章节"""
        doc.add_heading('3. 仿真结果', level=1)
        
        section_num = 1
        
        # 添加各类结果
        for fig_info in self.figures:
            doc.add_heading(f'3.{section_num} {fig_info["caption"]}', level=2)
            
            if fig_info.get('description'):
                doc.add_paragraph(fig_info['description'])
            
            # 保存并插入图片
            img_path = self._save_figure(fig_info['figure'], fig_info['name'])
            doc.add_picture(img_path, width=Inches(6))
            
            # 图片标题
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f'图 3.{section_num}: {fig_info["caption"]}').italic = True
            
            section_num += 1
        
        # 添加表格
        for table_info in self.tables:
            doc.add_heading(f'3.{section_num} {table_info["caption"]}', level=2)
            
            df = table_info['data']
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Light Shading Accent 1'
            
            # 添加表头
            for i, col in enumerate(df.columns):
                table.cell(0, i).text = str(col)
            
            # 添加数据
            for i, row in df.iterrows():
                for j, value in enumerate(row):
                    table.cell(i + 1, j).text = str(value)
            
            # 表格标题
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f'表 3.{section_num}: {table_info["caption"]}').italic = True
            
            section_num += 1
    
    def _add_conclusions_section(self, doc: Document):
        """添加结论章节"""
        doc.add_heading('4. 结论', level=1)
        
        if self.conclusions:
            for i, conclusion in enumerate(self.conclusions):
                doc.add_paragraph(f'{i+1}. {conclusion}', style='List Number')
        else:
            doc.add_paragraph('基于仿真结果，得出以下结论：')
            doc.add_paragraph('1. 材料相变行为符合预期', style='List Number')
            doc.add_paragraph('2. 机械性能满足设计要求', style='List Number')
            doc.add_paragraph('3. 建议的热处理工艺参数合理', style='List Number')
    
    def generate_word_report(self, output_path: str):
        """生成Word报告"""
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 创建报告各部分
        self._create_cover_page(doc)
        self._create_abstract(doc)
        self._add_composition_section(doc)
        self._add_parameters_section(doc)
        self._add_results_section(doc)
        self._add_conclusions_section(doc)
        
        # 添加页眉页脚
        self._add_header_footer(doc)
        
        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"Word报告已生成: {output_path}")
    
    def generate_excel_report(self, output_path: str):
        """生成Excel数据报告"""
        wb = Workbook()
        
        # 删除默认sheet
        wb.remove(wb.active)
        
        # 样式定义
        header_font = Font(name='微软雅黑', size=12, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        data_font = Font(name='微软雅黑', size=10)
        data_alignment = Alignment(horizontal='center', vertical='center')
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 1. 基本信息表
        ws_info = wb.create_sheet('基本信息')
        info_data = [
            ['项目', '内容'],
            ['报告标题', self.basic_info['title']],
            ['材料名称', self.material_info.get('name', '')],
            ['报告编号', self.basic_info['report_number']],
            ['作者', self.basic_info['author']],
            ['日期', self.basic_info['date']],
            ['软件', self.simulation_params['software']]
        ]
        
        for row in info_data:
            ws_info.append(row)
        
        # 应用样式
        for row in ws_info.iter_rows():
            for cell in row:
                cell.font = data_font
                cell.alignment = data_alignment
                cell.border = border
                if cell.row == 1:
                    cell.font = header_font
                    cell.fill = header_fill
        
        ws_info.column_dimensions['A'].width = 20
        ws_info.column_dimensions['B'].width = 40
        
        # 2. 成分表
        if self.material_info['composition']:
            ws_comp = wb.create_sheet('化学成分')
            
            # 表头
            headers = ['元素'] + list(self.material_info['composition'].keys())
            ws_comp.append(headers)
            
            # 数据
            values = ['含量 (wt%)'] + [f'{v:.3f}' for v in self.material_info['composition'].values()]
            ws_comp.append(values)
            
            # 样式
            for row in ws_comp.iter_rows():
                for cell in row:
                    cell.font = data_font
                    cell.alignment = data_alignment
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
            
            for col in ws_comp.columns:
                ws_comp.column_dimensions[col[0].column_letter].width = 12
        
        # 3. 相图数据
        if self.results.get('phase_diagram'):
            ws_phase = wb.create_sheet('相图数据')
            
            # 准备数据
            temps = self.results['phase_diagram']['temperatures']
            phases = self.results['phase_diagram']['phases']
            
            # 创建数据框架
            df_phase = pd.DataFrame(phases, index=temps)
            df_phase.index.name = '温度 (°C)'
            
            # 写入Excel
            for r_idx, row in enumerate([df_phase.index.name] + list(df_phase.columns)):
                ws_phase.cell(row=1, column=r_idx + 1, value=row)
            
            for r_idx, (temp, row) in enumerate(df_phase.iterrows()):
                ws_phase.cell(row=r_idx + 2, column=1, value=temp)
                for c_idx, value in enumerate(row):
                    ws_phase.cell(row=r_idx + 2, column=c_idx + 2, value=value)
            
            # 样式
            for row in ws_phase.iter_rows():
                for cell in row:
                    cell.font = data_font
                    cell.alignment = data_alignment
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
            
            # 添加图表
            if len(temps) > 1:
                chart = LineChart()
                chart.title = "平衡相图"
                chart.x_axis.title = "温度 (°C)"
                chart.y_axis.title = "相分数"
                
                data = Reference(ws_phase, min_col=2, min_row=1, 
                               max_col=len(phases) + 1, max_row=len(temps) + 1)
                cats = Reference(ws_phase, min_col=1, min_row=2, max_row=len(temps) + 1)
                
                chart.add_data(data, titles_from_data=True)
                chart.set_categories(cats)
                
                ws_phase.add_chart(chart, "H2")
        
        # 4. 机械性能数据
        if self.results.get('mechanical_properties'):
            ws_mech = wb.create_sheet('机械性能')
            
            mech_props = self.results['mechanical_properties']
            df_mech = pd.DataFrame(mech_props)
            
            # 写入数据
            for c_idx, col in enumerate(df_mech.columns):
                ws_mech.cell(row=1, column=c_idx + 1, value=col)
                for r_idx, value in enumerate(df_mech[col]):
                    ws_mech.cell(row=r_idx + 2, column=c_idx + 1, value=value)
            
            # 样式
            for row in ws_mech.iter_rows():
                for cell in row:
                    cell.font = data_font
                    cell.alignment = data_alignment
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
        
        # 5. 自定义表格
        for table_info in self.tables:
            ws_custom = wb.create_sheet(table_info['name'][:31])  # Excel sheet名称限制
            df = table_info['data']
            
            # 写入数据
            for c_idx, col in enumerate(df.columns):
                ws_custom.cell(row=1, column=c_idx + 1, value=str(col))
                for r_idx, value in enumerate(df[col]):
                    ws_custom.cell(row=r_idx + 2, column=c_idx + 1, value=value)
            
            # 样式
            for row in ws_custom.iter_rows():
                for cell in row:
                    cell.font = data_font
                    cell.alignment = data_alignment
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
        
        # 保存文件
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        wb.save(output_path)
        print(f"Excel报告已生成: {output_path}")
    
    def cleanup(self):
        """清理临时文件"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)


# 使用示例函数
def example_usage():
    """演示如何使用报告生成器"""
    # 创建报告生成器
    report = SimulationReportGenerator()
    
    # 设置基本信息
    report.set_basic_info(
        title="45号钢材料性能仿真分析报告",
        subtitle="基于JMatPro的综合分析",
        author="张三",
        company="XX材料科技有限公司"
    )
    
    # 设置材料信息
    report.set_material_info(
        name="45号钢",
        type="中碳钢",
        standard="GB/T 699-2015",
        description="常用的中碳结构钢，具有良好的综合机械性能"
    )
    
    # 添加成分
    composition = {
        'C': 0.45,
        'Si': 0.25,
        'Mn': 0.65,
        'Cr': 0.25,
        'Ni': 0.30,
        'P': 0.035,
        'S': 0.035
    }
    report.add_composition(composition)
    
    # 设置仿真参数
    report.set_simulation_params(
        software="JMatPro v12.0",
        modules=["相平衡", "TTT/CCT", "凝固", "机械性能"],
        temperature_range=[200, 1200],
        cooling_rates=[0.1, 1, 10, 100],
        grain_size=50,
        other_params={
            "奥氏体化温度": "900 °C",
            "保温时间": "30 min"
        }
    )
    
    # 添加相图结果（示例数据）
    temperatures = np.linspace(200, 1200, 100)
    phases = {
        'Ferrite': np.exp(-temperatures/500) * 0.8,
        'Austenite': 1 - np.exp(-temperatures/500) * 0.8,
        'Cementite': np.ones_like(temperatures) * 0.05
    }
    report.add_phase_diagram_results(temperatures.tolist(), phases)
    
    # 添加TTT结果（示例数据）
    ttt_data = {
        'Pearlite': {
            'start_times': [1, 5, 20, 100, 500],
            'start_temperatures': [700, 650, 600, 550, 500],
            'finish_times': [10, 50, 200, 1000, 5000],
            'finish_temperatures': [700, 650, 600, 550, 500]
        },
        'Bainite': {
            'start_times': [0.5, 2, 10, 50],
            'start_temperatures': [450, 400, 350, 300],
            'finish_times': [5, 20, 100, 500],
            'finish_temperatures': [450, 400, 350, 300]
        }
    }
    report.add_ttt_results(ttt_data)
    
    # 添加机械性能（示例数据）
    mech_props = {
        'temperature': [20, 100, 200, 300, 400, 500, 600],
        'yield_strength': [355, 340, 320, 290, 250, 180, 120],
        'tensile_strength': [600, 580, 550, 500, 430, 300, 200],
        'hardness': [180, 175, 170, 160, 145, 120, 90]
    }
    report.add_mechanical_properties(mech_props)
    
    # 添加凝固结果（示例数据）
    solidification = {
        'temperature': np.linspace(1500, 1300, 50).tolist(),
        'solid_fraction': np.linspace(0, 1, 50).tolist(),
        'liquidus_temp': 1480,
        'solidus_temp': 1420
    }
    report.add_solidification_results(solidification)
    
    # 添加自定义表格
    df_custom = pd.DataFrame({
        '热处理工艺': ['正火', '调质', '退火'],
        '温度 (°C)': [870, 840, 650],
        '保温时间 (min)': [60, 45, 120],
        '冷却方式': ['空冷', '油冷', '炉冷'],
        '硬度 (HB)': [180, 240, 160]
    })
    report.add_custom_table(df_custom, 'heat_treatment', '推荐热处理工艺参数')
    
    # 添加结论
    report.add_conclusion("45号钢在900°C奥氏体化后，采用不同冷却速率可获得不同的组织")
    report.add_conclusion("冷却速率在1-10°C/s范围内，主要获得珠光体+铁素体组织")
    report.add_conclusion("快速冷却(>100°C/s)可获得马氏体组织，硬度显著提高")
    report.add_conclusion("建议采用调质处理获得最佳的综合机械性能")
    
    # 生成报告
    report.generate_word_report("output/45钢仿真报告.docx")
    report.generate_excel_report("output/45钢仿真数据.xlsx")
    
    # 清理临时文件
    report.cleanup()


if __name__ == "__main__":
    example_usage()
